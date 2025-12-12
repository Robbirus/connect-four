from docx import Document
from datetime import datetime
import json

def generate_report(rules_text, training_summary, metrics):
    doc = Document()
    doc.add_heading('Connect Four - Project Report', level=1)
    doc.add_paragraph(f'Date: {datetime.utcnow().isoformat()}')
    doc.add_heading('Rules', level=2)
    doc.add_paragraph(rules_text)
    doc.add_heading('Dataset & Database', level=2)
    doc.add_paragraph(training_summary)
    doc.add_heading('Results', level=2)
    for k,v in metrics.items():
        doc.add_paragraph(f'{k}: {v}')
    doc.save('report.docx')

if __name__ == "__main__":
    rules = "Connect Four: two players alternate dropping pieces..."
    train_sum = "5000 games generated; 200k moves stored in SQLite"
    metrics = {'Accuracy': '45%', 'Avg Win Rate vs Random': '78%'}
    generate_report(rules, train_sum, metrics)
